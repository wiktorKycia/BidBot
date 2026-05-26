import logging
import os
import platform
import subprocess
import time
from pathlib import Path
from typing import Literal

import backoff
import docx
import fitz
import pandas as pd
from dotenv import load_dotenv
from google import genai
from google.genai.errors import ClientError
from pydantic import BaseModel, Field

from etl.scrapers.settings import ATTACHMENTS_DIR, BASE_DIR
from etl.loggers import setup_logging

load_dotenv(BASE_DIR / ".env")
setup_logging()
logger = logging.getLogger(__name__)

MissingDataFallback = Literal["unknown", "not_found", "needs_manual_review"]


class TenderData(BaseModel):
    tytul: str = Field(..., description="Pełny tytuł przetargu.")
    zamawiajacy: str | MissingDataFallback = Field(default="not_found", description="Nazwa podmiotu ogłaszającego przetarg.")
    budzet: str | MissingDataFallback = Field(
        default="not_found", description="Kwota budżetu wraz z walutą. Jeśli brak wprost w tekście, zwróć 'not_found'."
    )
    deadline: str | MissingDataFallback = Field(default="not_found", description="Ostateczny termin składania ofert.")
    miejsce_realizacji: str | MissingDataFallback = Field(default="not_found", description="Miejsce realizacji zamówienia.")
    wymagania_techniczne: list[str] = Field(default_factory=list, description="Lista kluczowych wymagań technicznych. Zostaw pustą listę, jeśli brak.")
    kryteria_oceny: list[str] = Field(default_factory=list, description="Kryteria oceny ofert.")
    wymagane_dokumenty: list[str] = Field(default_factory=list, description="Lista załączników i oświadczeń wymaganych od wykonawcy.")
    ryzyka: list[str] = Field(default_factory=list, description="Wymienione kary umowne, ryzyka lub szczególne obostrzenia w umowie.")


def extract_text_from_pdf(file_path: Path) -> str:
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        logger.error(f"Błąd PDF {file_path.name}: {e}")
    return text


def extract_text_from_docx(file_path: Path) -> str:
    try:
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Błąd DOCX {file_path.name}: {e}")
        return ""


def extract_data_from_excel(file_path: Path) -> str:
    try:
        engine = "xlrd" if file_path.suffix.lower() == ".xls" else "openpyxl"
        excel_file = pd.ExcelFile(file_path, engine=engine)
        all_sheets_text = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, engine=engine)
            if not df.empty:
                all_sheets_text.append(f"--- Arkusz: {sheet_name} ---")
                all_sheets_text.append(df.to_markdown(index=False))
        return "\n\n".join(all_sheets_text)
    except Exception as e:
        logger.error(f"Błąd Excel {file_path.name}: {e}")
        return ""


def extract_text_from_doc_windows(file_path: Path) -> str:
    try:
        import win32com.client
    except ImportError:
        logger.error("Brak biblioteki pywin32. Zainstaluj ją, aby czytać .doc na Windowsie.")
        return ""

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(file_path.resolve()), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
        return text
    except Exception as e:
        logger.error(f"Błąd Word COM {file_path.name}: {e}")
        return ""
    finally:
        if word:
            word.Quit()


def extract_text_from_doc_linux(file_path: Path) -> str:
    try:
        result = subprocess.run(["antiword", str(file_path)], capture_output=True, text=True)
        return result.stdout
    except FileNotFoundError:
        logger.error("Brak 'antiword' w systemie. Zainstaluj (np. apt-get install antiword).")
        return ""
    except Exception as e:
        logger.error(f"Błąd antiword {file_path.name}: {e}")
        return ""


def extract_text_from_doc(file_path: Path) -> str:
    if platform.system() == "Windows":
        return extract_text_from_doc_windows(file_path)
    else:
        return extract_text_from_doc_linux(file_path)


def parse_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        return extract_text_from_doc(file_path)
    elif suffix in [".xlsx", ".xls"]:
        return extract_data_from_excel(file_path)
    else:
        return ""


SYSTEM_PROMPT = """
Jesteś systemem eksperckim ds. zamówień publicznych.
Twoim zadaniem jest wyciągnięcie danych z załączonych dokumentów przetargowych.

REGUŁY BRAKU DANYCH:
1. Brak informacji wprost = "not_found"
2. Informacja jest, ale niejednoznaczna = "unknown"
3. Sprzeczne informacje w różnych załącznikach = "needs_manual_review"
4. ZERO halucynacji. Nie zgaduj dat, nie przeliczaj walut w pamięci.
"""

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))


def is_rate_limit(e: Exception) -> bool:
    if isinstance(e, ClientError) and e.code == 429:
        logger.warning("Limit API 429 osiągnięty, czekam przed wznowieniem prób...")
        return True
    return False


@backoff.on_exception(backoff.fibo, ClientError, max_tries=6, giveup=lambda e: not is_rate_limit(e), logger=logger)
def extract_tender_data(context_text: str) -> TenderData | None:
    model_name = "gemini-2.5-flash"

    response = client.models.generate_content(
        model=model_name,
        contents=f"Wyciągnij dane z poniższych załączników przetargowych:\n\n{context_text}",
        config={
            "system_instruction": SYSTEM_PROMPT,
            "temperature": 0.0,
            "response_mime_type": "application/json",
            "response_schema": TenderData,
        },
    )
    return response.parsed


def main(base_dir: Path = ATTACHMENTS_DIR):
    # Initialize database tables
    from etl.postgres_saver import create_tables
    try:
        create_tables()
    except Exception as db_err:
        logger.error(f"PostgreSQL connection/table creation error: {db_err}")

    base_path = Path(base_dir)

    if not base_path.exists():
        logger.error(f"Katalog {base_path} nie istnieje!")
        return

    for tender_dir in base_path.iterdir():
        if tender_dir.is_dir():
            logger.info(f"Rozpoczynam przetwarzanie przetargu: {tender_dir.name}")
            combined_text = []

            for file_path in tender_dir.glob("*"):
                if file_path.is_file() and not file_path.name.startswith("."):
                    logger.debug(f"Parsowanie pliku: {file_path.name}")
                    text = parse_file(file_path)
                    if text.strip():
                        combined_text.append(f"=== ZAŁĄCZNIK: {file_path.name} ===\n{text}")

            full_context = "\n\n".join(combined_text)

            if not full_context:
                logger.warning(f"Brak wyciągniętego tekstu w folderze {tender_dir.name}, pomijam.")
                continue

            logger.info(f"Wysyłam zapytanie LLM dla {tender_dir.name}...")
            try:
                tender_json_obj = extract_tender_data(full_context)
                if tender_json_obj:
                    logger.info(f"Pomyślnie wyciągnięto dane dla przetargu: {tender_dir.name}")
                    logger.info(f"WYNIK - JSON ({tender_dir.name}):\n{tender_json_obj.model_dump_json(indent=2)}")
                    
                    # Save extracted structured data to PostgreSQL
                    from etl.postgres_saver import save_llm_data
                    if save_llm_data(tender_dir.name, tender_json_obj):
                        logger.info(f"Pomyślnie zapisano dane LLM do bazy PostgreSQL dla przetargu: {tender_dir.name}")
                    else:
                        logger.warning(f"Nie udało się zapisać danych LLM do bazy PostgreSQL dla przetargu: {tender_dir.name}")
            except Exception as e:
                logger.error(f"Ostateczny błąd LLM po powtórkach dla {tender_dir.name}: {e}", exc_info=True)

            logger.info("Cooldown 3 sekundy")
            time.sleep(3)


if __name__ == "__main__":
    main()
